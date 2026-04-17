"""Multi-format document ingestion for UC3 podcast generation.

Normalizes PPTX / PDF / DOCX / TXT / MD / URL inputs into a `Document`
that downstream services (script generation, composition) consume.
"""

from __future__ import annotations

import logging
import mimetypes
import re
import uuid
from pathlib import Path
from urllib.parse import urlparse

from services.podcast_models import Document, Section

log = logging.getLogger(__name__)


class _SectionDeleted:
    pass


# Document/Section are imported from podcast_models.


# ---------------------------------------------------------------------------
# Format-specific ingesters
# ---------------------------------------------------------------------------

def ingest_pptx(path: str | Path) -> Document:
    """Parse a PPTX, render slide images via LibreOffice, return a Document."""
    from services.pptx_parser import parse_pptx

    p = Path(path)
    data = parse_pptx(p.read_bytes(), p.name)

    sections: list[Section] = []
    slide_titles: list[str] = []
    slide_notes: list[str] = []
    for sl in data.slides:
        heading = sl.title or f"Slide {sl.index + 1}"
        body = "\n".join(part for part in (sl.body, sl.notes) if part)
        sections.append(Section(heading=heading, text=body))
        slide_titles.append(sl.title or "")
        slide_notes.append(sl.notes or "")

    # Persist rendered PNGs alongside the source so compose can reference them.
    img_dir = p.parent / f"{p.stem}_slides_{data.id[:8]}"
    img_dir.mkdir(parents=True, exist_ok=True)
    slide_images: list[str] = []
    for idx, png in data.slide_images:
        out = img_dir / f"slide_{idx:03d}.png"
        out.write_bytes(png)
        slide_images.append(str(out))

    return Document(
        id=data.id,
        title=p.stem,
        sections=sections,
        slide_images=slide_images,
        slide_titles=slide_titles,
        slide_notes=slide_notes,
        source_kind="pptx",
    )


def ingest_pdf(path: str | Path) -> Document:
    """One section per PDF page."""
    from pypdf import PdfReader

    p = Path(path)
    reader = PdfReader(str(p))
    sections: list[Section] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            log.warning("pdf page %d extract failed: %s", i, exc)
            text = ""
        sections.append(Section(heading=f"Page {i + 1}", text=text.strip()))

    title = (reader.metadata.title if reader.metadata and reader.metadata.title else p.stem) or p.stem
    return Document(id=str(uuid.uuid4()), title=str(title), sections=sections, source_kind="pdf")


def ingest_docx(path: str | Path) -> Document:
    """Treat each Heading paragraph as a section boundary."""
    from docx import Document as DocxDocument

    p = Path(path)
    doc = DocxDocument(str(p))

    sections: list[Section] = []
    current = Section(heading=p.stem, text="")
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading"):
            if current.text or current.heading != p.stem:
                sections.append(current)
            current = Section(heading=text, text="")
        else:
            current.text = f"{current.text}\n{text}".strip() if current.text else text
    if current.text or not sections:
        sections.append(current)

    return Document(id=str(uuid.uuid4()), title=p.stem, sections=sections, source_kind="docx")


def ingest_text(path: str | Path) -> Document:
    """Plain text or markdown — split on top-level markdown headings if present."""
    p = Path(path)
    raw = p.read_text(encoding="utf-8", errors="replace")
    kind = "md" if p.suffix.lower() in {".md", ".markdown"} else "txt"

    sections: list[Section] = []
    if kind == "md" and re.search(r"^#{1,6}\s", raw, flags=re.MULTILINE):
        chunks = re.split(r"(?m)^(?=#{1,6}\s)", raw)
        for chunk in chunks:
            chunk = chunk.strip()
            if not chunk:
                continue
            lines = chunk.splitlines()
            heading = lines[0].lstrip("# ").strip()
            body = "\n".join(lines[1:]).strip()
            sections.append(Section(heading=heading, text=body))
    else:
        sections.append(Section(heading=p.stem, text=raw.strip()))

    return Document(id=str(uuid.uuid4()), title=p.stem, sections=sections, source_kind=kind)


def ingest_url(url: str) -> Document:
    """Fetch a URL and extract the main article body via readability."""
    import httpx
    from readability import Document as ReadabilityDocument
    from html import unescape

    resp = httpx.get(url, timeout=30, follow_redirects=True, headers={
        "User-Agent": "Mozilla/5.0 (compatible; ai-presenter/1.0)"
    })
    resp.raise_for_status()

    rd = ReadabilityDocument(resp.text)
    title = rd.short_title() or urlparse(url).netloc
    summary_html = rd.summary(html_partial=True)
    text = re.sub(r"<[^>]+>", " ", summary_html)
    text = unescape(re.sub(r"\s+", " ", text)).strip()

    return Document(
        id=str(uuid.uuid4()),
        title=title,
        sections=[Section(heading=title, text=text)],
        source_kind="url",
    )


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def ingest(path_or_url: str | Path) -> Document:
    """Pick the right ingester from extension or URL scheme."""
    s = str(path_or_url)
    parsed = urlparse(s)
    if parsed.scheme in {"http", "https"}:
        return ingest_url(s)

    p = Path(s)
    ext = p.suffix.lower()
    if ext == ".pptx":
        return ingest_pptx(p)
    if ext == ".pdf":
        return ingest_pdf(p)
    if ext == ".docx":
        return ingest_docx(p)
    if ext in {".md", ".markdown", ".txt"}:
        return ingest_text(p)

    # Fallback: sniff content type
    ctype, _ = mimetypes.guess_type(s)
    if ctype:
        if "pdf" in ctype:
            return ingest_pdf(p)
        if "word" in ctype or "officedocument" in ctype:
            return ingest_docx(p)
    return ingest_text(p)
